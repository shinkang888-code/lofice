#!/usr/bin/env python3
"""office-convert API — Phase 3 LibreOffice headless + HWPX→DOCX (CM 변환)"""

from __future__ import annotations

import base64
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Optional

import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="lofice office-convert API", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SOFFICE = os.environ.get("LIBREOFFICE_PATH", "soffice")


def has_libreoffice() -> bool:
    return bool(shutil.which(SOFFICE))


def run_soffice_convert(src: Path, out_dir: Path, target: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    proc = subprocess.run(
        [
            SOFFICE,
            "--headless",
            "--norestore",
            "--convert-to",
            target,
            "--outdir",
            str(out_dir),
            str(src),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if proc.returncode != 0:
        raise HTTPException(status_code=500, detail=(proc.stderr or proc.stdout or "libreoffice failed")[:2000])
    matches = list(out_dir.glob(f"*.{target}"))
    if not matches:
        raise HTTPException(status_code=500, detail="conversion produced no output")
    return matches[0]


def hwpx_to_docx_simple(hwpx_path: Path, out_path: Path) -> None:
    """경량 HWPX→DOCX — 텍스트 추출 후 python-docx (LibreOffice 없을 때)"""
    import zipfile
    import xml.etree.ElementTree as ET

    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx not installed")

    texts: list[str] = []
    with zipfile.ZipFile(hwpx_path, "r") as zf:
        for name in sorted(zf.namelist()):
            if name.endswith(".xml") and "section" in name.lower():
                root = ET.fromstring(zf.read(name))
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        texts.append(elem.text.strip())
                    if elem.tail and elem.tail.strip():
                        texts.append(elem.tail.strip())

    doc = Document()
    doc.add_heading(hwpx_path.stem, level=1)
    for t in texts[:5000]:
        doc.add_paragraph(t)
    doc.save(str(out_path))


LEGACY_TO_OOXML = {
    ".doc": "docx",
    ".dot": "docx",
    ".xls": "xlsx",
    ".xlt": "xlsx",
    ".ppt": "pptx",
    ".pps": "pptx",
    ".pot": "pptx",
}


def target_format(suffix: str) -> str | None:
    return LEGACY_TO_OOXML.get(suffix.lower())


@app.post("/normalize")
async def normalize_office(file: UploadFile = File(...)) -> dict[str, str]:
    """레거시 Office → OOXML 정규화 (Phase 1)"""
    suffix = Path(file.filename or "doc").suffix.lower()
    target = target_format(suffix)
    if suffix in {".docx", ".xlsx", ".pptx", ".docm", ".xlsm", ".pptm"}:
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="empty file")
        kind = suffix.lstrip(".")
        if kind.endswith("m"):
            kind = kind[:-1] + "x"
        return {
            "file_name": Path(file.filename or "document").stem + f".{kind}",
            "data_base64": base64.b64encode(data).decode(),
            "format": kind,
        }
    if not target:
        raise HTTPException(status_code=400, detail="unsupported format for normalize")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="empty file")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        src = tmp_path / f"input{suffix}"
        src.write_bytes(data)
        if has_libreoffice():
            out = run_soffice_convert(src, tmp_path, target)
        else:
            raise HTTPException(
                status_code=503,
                detail="LibreOffice required for legacy Office normalization",
            )
        return {
            "file_name": Path(file.filename or "document").stem + f".{target}",
            "data_base64": base64.b64encode(out.read_bytes()).decode(),
            "format": target,
        }


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "libreoffice": has_libreoffice(),
        "soffice": SOFFICE,
        "features": ["hwpx-to-docx", "libreoffice-convert", "normalize"],
    }


@app.post("/convert/hwpx-to-docx")
async def convert_hwpx_to_docx(file: UploadFile = File(...)) -> dict[str, str]:
    suffix = Path(file.filename or "doc.hwpx").suffix.lower()
    if suffix != ".hwpx":
        raise HTTPException(status_code=400, detail="only .hwpx supported")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="empty file")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        src = tmp_path / "input.hwpx"
        out = tmp_path / "output.docx"
        src.write_bytes(data)

        if has_libreoffice():
            run_soffice_convert(src, tmp_path, "docx")
            out = next(tmp_path.glob("*.docx"))
        else:
            hwpx_to_docx_simple(src, out)

        return {
            "file_name": Path(file.filename or "document.hwpx").stem + ".docx",
            "data_base64": base64.b64encode(out.read_bytes()).decode(),
            "format": "docx",
        }


@app.post("/convert/libreoffice")
async def convert_libreoffice(
    file: UploadFile = File(...),
    target: str = Form("docx"),
) -> dict[str, str]:
    if not has_libreoffice():
        raise HTTPException(status_code=503, detail="LibreOffice (soffice) not available on server")
    if target not in {"docx", "pdf", "odt", "html", "xlsx", "pptx"}:
        raise HTTPException(status_code=400, detail="target must be docx|xlsx|pptx|pdf|odt|html")

    data = await file.read()
    suffix = Path(file.filename or "document").suffix or ".bin"
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        src = tmp_path / f"input{suffix}"
        src.write_bytes(data)
        out = run_soffice_convert(src, tmp_path, target)
        return {
            "file_name": out.name,
            "data_base64": base64.b64encode(out.read_bytes()).decode(),
            "format": target,
        }


def main() -> None:
    host = os.environ.get("OFFICE_CONVERT_HOST", "0.0.0.0")
    # Render/Heroku inject PORT; local default 8200
    port = int(os.environ.get("PORT", os.environ.get("OFFICE_CONVERT_PORT", "8200")))
    print(f"office-convert API on http://{host}:{port} (libreoffice={has_libreoffice()})")
    uvicorn.run(app, host=host, port=port)


if __name__ == "__main__":
    main()
